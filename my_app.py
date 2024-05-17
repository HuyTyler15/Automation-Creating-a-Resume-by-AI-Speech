from docx import Document
from docx.shared import Inches
import pyttsx3

pyttsx3.speak('Hi guys')

def speak(text):
    pyttsx3.speak(text)

document = Document()

document.add_picture(
    'tyler.jpg', 
    width=Inches(1.2))

# name and mail
my_name = input('What is your name ? ')
speak('Hello '+ name + 'How are you today ?')
speak('What is your number ? ')
phone = input('What is your number ? ')
speak('and your email ? ')
mail = input('and your email ? ')
speak('Let me know your skype name. ')
skype = input ('Let me know your skype name. ')

n = document.add_heading( my_name, 0)
document.add_paragraph('Phone: '+ phone + ' | Mail: '+ mail + ' | Skype: '+ skype)

# about me

document.add_heading('ABOUT ME')
speak('Tell me about yourself ? ')
about_me = input('Tell me about yourself ? ')
document.add_paragraph(about_me)

# skills

speak('What is your skill and please show it all to us ')
document.add_heading('SKILLS')
speak('What is your skill ? ')
skill = input('What is your skill ? ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    speak('Do you have any skills to show up? Yes or No ')
    has_more_skills = input('Do you have any skills to show up? Yes or No ')
    if has_more_skills.lower() == 'yes':
        speak('What is that ? Detail, please. ')
        skill = input('What is that ? Detail, please. ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

# education
speak('Let me know about your education ')

document.add_heading('EDUCATION')
speak('Which university do you study? ')
university = input('Which university do you study? ')
u = document.add_paragraph()
speak('What is your major at the' + university +' ? ')
major = input('What is your major at the' + university +' ? ')
speak('What is a Bachelor degree ? ')
bachelor = input('What is a Bachelor degree ? ')
speak('From date')
start = input('From date : ')
speak('To date')
end = input('To date: ')

u.add_run('The university ' + university + ' | ').bold = True
u.add_run(start + '-' + end + '\n').italic
u.add_run(bachelor + ' : '+ major + ' .')

while True:
    speak('You have another one? Yes or No ')
    has_more_educations = input('You have another one? Yes or No ')
    if has_more_educations.lower() == 'yes':
        speak('What is the fullname of your university or college? ')
        university = input('What is the fullname of your university or college? ')
        u = document.add_paragraph()
        speak('What is your major at the' + university +' ? ')
        major = input('What is your major at the' + university +' ? ')
        speak('What is a Bachelor degree ? ')
        bachelor = input('What is a degree ? ')
        start = input('From date : ')
        speak('To date')
        end = input('To date: ')

        u.add_run(university + ' ').bold = True
        u.add_run(start + '-' + end + '\n').italic
        u.add_run(bachelor + ' : '+ major)
        
    else:
        break

# work experiences

speak('How many company you used to work before ?')
document.add_heading('WORK EXPERIENCE')
p = document.add_paragraph()
speak('What is your lastest company ? ')
company = input('What is your lastest company ? ')
speak('What was the position in' + company+ ' ? ')
position = input('What was the position in' + company+ ' ? ')
speak('From date : ')
from_date = input('From date : ')
speak('To date: ')
to_date = input('To date: ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic
p.add_run(position + '\n').bold = True

speak( 'Let me know more about you, could you please describe about your experience in a ' 
                    + company + ' ? ')
experience_details = input( 'Let me know more about you, could you please describe about your experience in a ' 
                    + company + ' ? ')
p.add_run(experience_details + '\n')

speak('Anything else ? ')
anything_else = input('Anything else ? ')
p.add_run(anything_else)

# more experiences

while True:
    speak(
        'Do you have more company in the past ? Yes or No '
    )
    has_more_experience = input(
        'Do you have more company in the past ? Yes or No '
    )
    if has_more_experience.lower() == 'yes':
        p = document.add_paragraph()

        speak('What name of that company ? ')
        company = input('What name of that company ? ')
        speak('What was the position in' + company+ ' ? ')
        position = input('What was the position in' + company+ ' ? ')
        speak('From date : ')
        from_date = input('From date : ')
        speak('To date: ')
        to_date = input('To date: ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic
        p.add_run(position).bold = True

        speak( 'Describe your experiences in the ' 
                            + company + ' ? ')
        experience_details = input( 'Describe your experiences in the ' 
                            + company + ' ? ')

        p.add_run(experience_details + '/n')
        speak('Anything else ? ')
        anything_else = input('Anything else ? ')
        p.add_run(anything_else + '/n')
    else:
        break

# hobbies
speak('your hobbies will be fun ')
document.add_heading('HOBBIES')
speak('What is your hobby after working time ? ')
hobbie = input('What is your hobby after working time ? ')
h = document.add_paragraph(hobbie)
h.style = 'List Bullet'

while True:
    speak('Do you have any hobby ? Yes or No ')
    more_hobbies = input('Do you have any hobby ? Yes or No ')
    if more_hobbies.lower() == 'yes':
        speak('What is another hobby after working time ? ')
        hobbie = input('What is another hobby after working time ? ')
        h = document.add_paragraph(hobbie)
        h.style = 'List Bullet'
    else:
        break

# footer

section = document.sections[0]
footer = section.footer
footer_para = footer.paragraphs[0]
speak('This CV generated by huytyler15@gmail.com by using Python for coding')
footer_para.text = "This CV generated by huytyler15@gmail.com by using Python for coding "

document.save('CV_huytyler.docx')